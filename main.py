from docx import Document
from docx.shared import Inches

# Criando o documento
doc = Document()
doc.add_heading('Estimativa de Emissão de CO₂ pelos Veículos em São Caetano do Sul', 0)

# Adicionando as seções
doc.add_heading('Premissas', level=1)
doc.add_paragraph(
    '- Frota estimada em São Caetano do Sul: 110.000 veículos\n'
    '- Quilometragem média mensal por veículo: 1.200 km/mês\n'
    '- Consumo médio de combustível: 10 km/l\n'
    '- Emissão média de CO₂ por litro de gasolina: 2,3 kg CO₂/l\n'
    '(Fonte: EPA - Environmental Protection Agency)'
)

doc.add_heading('Cálculos', level=1)

doc.add_heading('1. Consumo mensal de combustível por veículo', level=2)
doc.add_paragraph('Litros por mês = 1.200 km ÷ 10 km/l = 120 litros/mês')

doc.add_heading('2. Emissão de CO₂ por veículo por mês', level=2)
doc.add_paragraph('CO₂ por veículo = 120 litros × 2,3 kg CO₂/litro = 276 kg CO₂/mês')

doc.add_heading('3. Emissão total da frota por mês', level=2)
doc.add_paragraph('Total CO₂ = 276 kg × 110.000 = 30.360.000 kg CO₂/mês = 30.360 toneladas de CO₂/mês')

doc.add_heading('Conclusão', level=1)
doc.add_paragraph(
    'A frota automotiva de São Caetano do Sul é responsável pela emissão de aproximadamente:\n'
    '30.360 toneladas de CO₂ por mês.\n'
    'Esse valor reforça a importância da adoção de políticas de mobilidade urbana sustentável e eletrificação veicular.'
)

doc.add_heading('Referências', level=1)
doc.add_paragraph(
    '- EPA – Environmental Protection Agency: https://www.epa.gov/greenvehicles\n'
    '- DENATRAN – Frota veicular brasileira: https://www.gov.br/infraestrutura\n'
    '- IPCC – Intergovernmental Panel on Climate Change, 2021'
)

doc.add_heading('Curiosidade', level=1)
doc.add_paragraph(
    'Um ônibus elétrico pode reduzir até 60 toneladas de CO₂ por ano em comparação com um modelo a diesel.\n'
    'Fonte: https://cleantechnica.com/2020/08/19/electric-buses-save-60-tons-of-co2-year'
)

# Salvando o documento
file_path = "/mnt/data/Estimativa_Emissao_CO2_SCS.docx"
doc.save(file_path)

file_path
